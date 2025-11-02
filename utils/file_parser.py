"""
File parser utility for handling PDF, DOCX, and TXT files.
"""

import io
import PyPDF2
import pdfplumber
from docx import Document
import streamlit as st
from typing import Optional

class FileParser:
    """Utility class for parsing different file formats."""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file using multiple methods for better accuracy.
        
        Args:
            file_content (bytes): PDF file content
            
        Returns:
            str: Extracted text
        """
        text = ""
        
        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If pdfplumber didn't extract much text, try PyPDF2
            if len(text.strip()) < 50:
                text = ""
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        
        except Exception as e:
            st.error(f"Error extracting text from PDF: {str(e)}")
            return ""
        
        return text.strip()
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content (bytes): DOCX file content
            
        Returns:
            str: Extracted text
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        
        except Exception as e:
            st.error(f"Error extracting text from DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_content (bytes): TXT file content
            
        Returns:
            str: Extracted text
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='ignore')
            return text.strip()
        
        except Exception as e:
            st.error(f"Error extracting text from TXT: {str(e)}")
            return ""
    
    @staticmethod
    def parse_uploaded_file(uploaded_file) -> Optional[str]:
        """
        Parse uploaded file and extract text based on file type.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Optional[str]: Extracted text or None if parsing failed
        """
        if uploaded_file is None:
            return None
        
        file_content = uploaded_file.read()
        file_type = uploaded_file.type
        file_name = uploaded_file.name.lower()
        
        # Determine file type and parse accordingly
        if file_type == "application/pdf" or file_name.endswith('.pdf'):
            return FileParser.extract_text_from_pdf(file_content)
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.endswith('.docx'):
            return FileParser.extract_text_from_docx(file_content)
        
        elif file_type == "text/plain" or file_name.endswith('.txt'):
            return FileParser.extract_text_from_txt(file_content)
        
        else:
            st.error(f"Unsupported file type: {file_type}. Please use PDF, DOCX, or TXT files.")
            return None
    
    @staticmethod
    def validate_extracted_text(text: str, min_length: int = 50) -> bool:
        """
        Validate if extracted text is meaningful.
        
        Args:
            text (str): Extracted text
            min_length (int): Minimum required text length
            
        Returns:
            bool: True if text is valid, False otherwise
        """
        if not text or len(text.strip()) < min_length:
            return False
        
        # Check if text contains mostly readable characters
        readable_chars = sum(1 for char in text if char.isalnum() or char.isspace() or char in '.,!?-()[]{}')
        total_chars = len(text)
        
        if total_chars == 0:
            return False
        
        readable_ratio = readable_chars / total_chars
        return readable_ratio > 0.7  # At least 70% readable characters
    
    @staticmethod
    def preview_text(text: str, max_length: int = 200) -> str:
        """
        Create a preview of extracted text.
        
        Args:
            text (str): Full text
            max_length (int): Maximum preview length
            
        Returns:
            str: Text preview
        """
        if not text:
            return "No text extracted"
        
        preview = text.strip()
        if len(preview) > max_length:
            preview = preview[:max_length] + "..."
        
        return preview